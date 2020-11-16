from django.shortcuts import render,redirect,HttpResponse,HttpResponseRedirect
from django.contrib import messages

from django.utils.encoding import smart_str
# Importing python classes
from genque.PythonClasses.HeaderFooterRemover import pdf2text_all, pdf2text_spec

# Importing packages for accuracy testing
from sklearn.metrics import confusion_matrix
from sklearn.metrics import average_precision_score
import numpy as np

# Importing packages for generate
import spacy
from gender_detector import gender_detector as gd
from nltk import *
from nltk.tokenize import RegexpTokenizer, MWETokenizer
from stop_words import get_stop_words
from nltk.stem.porter import PorterStemmer
from gensim import corpora, models
from docx import Document
from docx.shared import Inches
import gensim
import language_check
import spacy
import random
import datetime
import uuid

tnf_words = []
tnf_list = []
wordlist = []
nlp = spacy.load('en_core_web_sm')
from spacy.lang.en import English
def remove_whitespace_entities(doc):
    doc.ents = [e for e in doc.ents if not e.text.isspace()]
    return doc
nlp.add_pipe(remove_whitespace_entities, after='ner')
def distractor_maker(ans, entity, choices,distractor, iterator):

	while not len(choices) >= 3:
		n = random.randrange(0,len(entity))
		distractor = entity[n]
		
		if not ans == distractor and not distractor in choices:
			choices.append(distractor)

	return choices[2]

	# if len(choices) == 3:
	# 	pass
	# else:
	# 	n = random.randrange(0,len(entity))
	# 	distractor = entity[n]
	# 	if not ans == distractor:
	# 		choices.append(distractor)
			
	# 	return distractor_maker(ans, entity, choices, distractor, iterator)

def MCQ(text,ans,ORG,DATE,PERSON,GPE_LOCATION,EVENT):

	choices = []
	distractor = ans

	if ans == 'ORG':
		choices.append(text)
		if not len(ORG) < 3:
			choices.append(distractor_maker(text, ORG, choices,distractor,0))
		elif not len(PERSON) < 3:
			choices.append(distractor_maker(text, PERSON, choices,distractor,0))
		elif not len(GPE_LOCATION) < 3:
			choices.append(distractor_maker(text, GPE_LOCATION, choices,distractor,0))

	if ans == 'DATE':
		choices.append(text)
		if not len(DATE) < 3:
			choices.append(distractor_maker(text, DATE, choices,distractor,0))
		elif not len(EVENT) < 3:
			choices.append(distractor_maker(text, EVENT, choices,distractor,0))
		elif not len(ORG) < 3:
			choices.append(distractor_maker(text, ORG, choices,distractor,0))
	
	if ans == 'PERSON':
		choices.append(text)
		if not len(PERSON) < 3:
			choices.append(distractor_maker(text, PERSON, choices,distractor,0))
		elif not len(ORG) < 3:
			choices.append(distractor_maker(text, ORG, choices,distractor,0))
		elif not len(GPE_LOCATION) < 3:
			choices.append(distractor_maker(text, GPE_LOCATION, choices,distractor,0))
	
	if ans == 'GPE':
		choices.append(text)
		if not len(GPE_LOCATION) < 3:
			choices.append(distractor_maker(text, GPE_LOCATION, choices,distractor,0))
		elif not len(ORG) < 3:
			choices.append(distractor_maker(text, ORG, choices,distractor,0))
		elif not len(PERSON) < 3:
			choices.append(distractor_maker(text, PERSON, choices,distractor,0))

	if ans == 'EVENT':
		choices.append(text)
		if not len(EVENT) < 3:
			choices.append(distractor_maker(text, EVENT, choices,distractor,0))
		elif not len(DATE) < 3:
			choices.append(distractor_maker(text, DATE, choices,distractor,0))
		elif not len(ORG) < 3:
			choices.append(distractor_maker(text, ORG, choices,distractor,0))

	for idx,index in enumerate(choices):
		choices[idx] = str(index).capitalize()

	return choices[:3]


def trueOrFalsetemp(label,text,ORG,PERSON,GPE_LOCATION,DATE,EVENT):
	textMatcher = text

	if label == "ORG" and len(ORG) > 1:
		while text == textMatcher:
			n = random.randrange(0,len(ORG))
			textMatcher = ORG[n]
	else:
		label = "PERSON"

	if label == "PERSON" and len(PERSON) > 1:
		while text == textMatcher:
			n = random.randrange(0,len(PERSON))
			textMatcher = PERSON[n]
	else:
		label = "GPE"

	if label == "GPE" and len(GPE_LOCATION) > 1:
		while text == textMatcher:
			n = random.randrange(0,len(GPE_LOCATION))
			textMatcher = GPE_LOCATION[n]
	else:
		label = "DATE"

	if label == "DATE" and len(DATE) > 1:
		while text == textMatcher:
			n = random.randrange(0,len(DATE))
			textMatcher = DATE[n]
	else:
		label = "EVENT"

	if label == "EVENT" and len(EVENT) > 1:
		while text == textMatcher:
			n = random.randrange(0,len(EVENT))
			textMatcher = EVENT[n]
	else:
		label = "ORG"

	return textMatcher

def UniqueItems(Noun):
	seen = set()
	result = []
	for item in Noun:
		if item not in seen:
			seen.add(item)
			result.append(item)

	return result
def dataCleansing(Words):
    nlp = spacy.load('en')
    remove_space = re.sub(r'["\n]', r' ', Words)
    parse = nlp(remove_space)
    sentences = [sent.string.strip() for sent in parse.sents]
    
    return sentences

def remove_whitespace_entities(doc):
    doc.ents = [e for e in doc.ents if not e.text.isspace()]
    return doc
def UniqueItems(Noun):
    seen = set()
    result = []
    for item in Noun:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result

def getPrecisionRecall(instance):
    total_recall = 0
    total_precision = 0
    total_f_measure = 0
    true_positive = 0
    false_negative = 0
    fp = np.sum(instance,axis=0) 

    for i in range(len(instance)):
        
        true_positive = instance[i][i]
        false_negative = sum(instance[i])
        if not false_negative == 0:
            recall = true_positive/false_negative
        else:
            recall = 1.0
        false_positive = fp[i]
        
        if not false_positive == 0:
            precision = true_positive/false_positive
        else:
            precision = 1.0
            
        f_measure = 2 * ((precision*recall)/(precision+recall))
        
        total_precision += precision
        total_recall += recall
        total_f_measure += f_measure
        print("RECALL", recall, "PRECISION", precision, "F_MEASURE", f_measure)
        
    avg_recall = total_recall/len(instance)
    avg_precision = total_precision/len(instance)
    avg_f_measure = total_f_measure/len(instance)
    print("AVERAGE OF RECALL", avg_recall, "AVERAGE of PRECISION", avg_precision, "AVERAGE OF F_MEASURE", avg_f_measure)

    return [avg_recall,avg_precision,avg_f_measure]

# Created views
def index(request):
    return render(request, 'genque/home.html')


def sample(request):
    return render(request, 'genque/sample.html')


# noinspection PyUnreachableCode
def generate(request):
	before = datetime.datetime.now()

	# Process with PDF
	pdfFileObj = request.FILES['file']
	pageFrom = int(request.POST.get('pageFrom'))
	pageTo = int(request.POST.get('pageTo'))

	# Removing Header and Footer in PDF file
	pdfObj = pdf2text_spec(pdfFileObj,pageFrom,pageTo)

	# Stands as input text
	text = ""

	# Retrieving every pages in pdfObj and append into text varible
	for page in pdfObj:
	    page = re.sub(r'["\n]', r' ', page)  # Removing next line in text
	    dn = nlp(page)
	    page_sentences = [sent.string.strip() for sent in dn.sents]
	    for sentence in range(len(page_sentences)):

	    	if sentence < 3:

	    		pass

	    	elif sentence > len(page_sentences)-3:

	    		pass
	    		
	    	else:


	    		text += page_sentences[sentence].replace("- ","")

	texts = re.sub("[\(\[].*?[\)\]]", "", text)

	detector = gd.GenderDetector('us')  # It can also be ar, uk, uy.

	# Initiate list of label for NER
	ORG = []
	PERSON = []
	GPE_LOCATION = []
	DATE = []
	EVENT = []

	# Initiate list of question variable
	questions = []

	# ------------------------------------------------------PREPROCESSING---------------------------------------------------

	# Tokenized text by sentence to become document in every
	#lda_sentence_tokens = sent_tokenize(text)
	dm = nlp(texts, disable=['tagger', 'ner'])
	lda_sentences = [sent.string.strip() for sent in dm.sents]
	
	lda_sentence_tokens = []
	sentence_rule = ["nsubj","dobj","pobj"]
	for sentence in lda_sentences:
		doc = nlp(sentence)

		for chunk in doc.noun_chunks:
			if any(x == chunk.root.dep_ for x in sentence_rule):
				lda_sentence_tokens.append(sentence)
				break

	# for a in lda_sentences:
	# 	if not a in lda_sentence_tokens:
	# 		print(a)

	# Initiate tokenizer
	tokenizer = RegexpTokenizer(r'\w+')

	# Create English stop words list
	en_stop = get_stop_words('english')

	# Create p_stemmer of class PorterStemmer
	p_stemmer = PorterStemmer()

	# Compile sample documents into a list
	doc_set = lda_sentence_tokens

	# List for tokenized documents in loop
	texts = []
	text_cleanse = ""

	# Collection of NER
	
	doc = nlp(text)

	for ent in doc.ents:
	    
	    if not ent.text == ' ' and not ent.text == '' and not ent.text == '  ' :

	        if ent.label_ == 'ORG' and not ent.text == ' ' and not ent.text == '':

	            ORG.append(ent.text)

	        elif ent.label_ == 'PERSON' and not ent.text == ' ' and not ent.text == '':

	            PERSON.append(ent.text)

	        elif ent.label_ == 'GPE' and not ent.text == ' ' and not ent.text == '':
	            GPE_LOCATION.append(ent.text)

	        elif ent.label_ == 'DATE':
	            DATE.append(ent.text)

	        elif ent.label_ == 'EVENT':
	            EVENT.append(ent.text)


	print("ORG", len(ORG))
	print("PERSON", len(PERSON))
	print("GPE_LOCATION", len(GPE_LOCATION))
	print("DATE", len(DATE))
	print("EVENT", len(EVENT))
	NER = ORG + PERSON + GPE_LOCATION + DATE + EVENT
	print("THIS IS NER")
	print(NER)

	# Coop through document list
	for i in doc_set:

	    # Clean and tokenize document string
	    # Remove the NER collected words before tokenize to aviod tokenized the NER collection
	    for token in range(len(NER)):
	        if NER[token] in i:  # Replace the tokens to exclude the NER tokens
	            i = i.replace(NER[token], " tags" + str(token) + " ")

	    raw = i.lower()  #
	    tokens = tokenizer.tokenize(raw)

	    # Recall the words collected in NER
	    for word in range(len(tokens)):
	        for token in range(len(NER)):

	            if tokens[word] == ("tags" + str(token) + ""):
	                tokens[word] = NER[token]

	    # Remove stop words from tokens
	    stopped_tokens = [i for i in tokens if not i in en_stop]
	    questions.append(tokens)

	    # Add tokens to list
	    texts.append(stopped_tokens)

	# ---------------------------------------------------TOPIC MODELING (LDA)------------------------------------------------
	# Turn our tokenized documents into a id <-> term dictionary
	dictionary = corpora.Dictionary(texts)

	# Convert tokenized documents into a document-term matrix
	corpus = [dictionary.doc2bow(text) for text in texts]

	# Generate LDA model
	ldamodel = gensim.models.ldamulticore.LdaMulticore(corpus, num_topics=5, id2word=dictionary, workers = 2)
	lda_raw = ldamodel.print_topics(num_topics=5, num_words=100)

	# ------------------------------------------- POST-PROCESSING----------------------------------------------------------------
	# Data Cleansing
	lda = []
	for topic in range(len(lda_raw)):
	    nstr = re.sub(r'[""|,|\\|''|)|(]', r'', str(lda_raw[topic]))
	    lda.append(nstr)

	# separate by topic
	pre_tokens = []
	for x in lda:
	    pre_tokens.append(str(x).split(" + "))

	# separate by weights and word
	lda_tokens = []
	for i in pre_tokens:
	    for x in i:
	        lda_tokens.append(str(x).split('*'))

	lda_words = []
	index = 0
	for i in lda_tokens:
	    if index >= 2 and index < len(lda_tokens) - 2:
	        lda_words.append(i[1])
	    index += 1



	for weight in lda_tokens:
		lda_tokens.index(weight)
		for words in range(len(weight)):

			if words == 0:
				weight[words].replace(str(len(weight))+" '","")

	i = 0
	cleanse = ''
	sentsw = ''
	weights = 0
	num = 0
	count = 0
	sentences = ''
	sentsWeight = [0]* len(lda_sentence_tokens)
	for index in lda_tokens:

		cleanse = str(index[0]).replace(str(i),"")
		sentsw = cleanse.replace(" '","")
		weightWords = str(index[1]).replace(str(i),"") 

		for sents in range(len(lda_sentence_tokens)):

			if weightWords in lda_sentence_tokens[sents]:
				count = str(lda_sentence_tokens[sents]).count(weightWords)
				weights += float(sentsw) * count
				num = sentsWeight[sents]
				sentsWeight[sents] += float(weights)
				continue
	i += 1
	
	yx = zip(sentsWeight,lda_sentence_tokens)

	sentenceWeighted = sorted(yx, key=lambda x: x[0])

	for sents in range(len(sentenceWeighted)):
		sentsWeight[sents] = sentenceWeighted[sents][1]

	trueOrFalse = []
	NER_names = []
	NER_Cleanse = UniqueItems(NER)
	question_finb_key = []
	for name in NER_Cleanse:
		NER_names.append(re.sub(r'[^\w]', ' ', name))

	question_finb = []

	for sentence in sentsWeight:

		for j in range(20):
			if j > len(NER_names)-1:
				break
			if NER_names[j] in question_finb_key:

				continue

			elif NER_names[j] in sentence and j < len(NER_names) and NER_names[j] != '' and NER_names != ' ' and len(str(sentence)) >= 25:
				limit = 1
				temp_word = str(sentence).replace(NER_names[j], ' ___________ ',1)
				question_finb_key.append(NER_names[j])
				question_finb.append(temp_word)
				j = j
				limit = 0

				break

			

	question_finb_key_random = question_finb_key


	tnf_list = []
	tnf_str = ''

	tnf_temp = UniqueItems(NER_Cleanse)

	for tnf in tnf_temp:

		tnf_str += tnf+" "

	tnf_words = nlp(tnf_str)

	count = 0
	question_tf = []
	question_tf_key = []
	for sentence in sentsWeight:

	    for dep in tnf_words.ents:
	    	if dep.text in sentence:
			    if dep.label_ == 'ORG':
			    	temp_word = str(sentence).replace(dep.text, str(trueOrFalsetemp(dep.label_,dep.text,ORG,PERSON,GPE_LOCATION,DATE,EVENT)))
			    	question_tf.append(temp_word+"\n")
			    	question_tf_key.append(dep.text)
			    	count +=1
			    	break

			    if dep.label_ == 'PERSON':
			    	temp_word = str(sentence).replace(dep.text, str(trueOrFalsetemp(dep.label_,dep.text,ORG,PERSON,GPE_LOCATION,DATE,EVENT)))
			    	question_tf.append(temp_word)
			    	question_tf_key.append(dep.text)
			    	count +=1
			    	break

			    if dep.label_ == 'GPE':
			    	temp_word = str(sentence).replace(dep.text, str(trueOrFalsetemp(dep.label_,dep.text,ORG,PERSON,GPE_LOCATION,DATE,EVENT)))
			    	question_tf.append(temp_word)
			    	question_tf_key.append(dep.text)
			    	count+=1
			    	break

			    if dep.label_ == 'DATE':
			    	temp_word = str(sentence).replace(dep.text,str(trueOrFalsetemp(dep.label_,dep.text,ORG,PERSON,GPE_LOCATION,DATE,EVENT)))
			    	question_tf.append(temp_word)
			    	question_tf_key.append(dep.text)
			    	count+=1
			    	break

			    if dep.label_ == 'EVENT':
			    	temp_word = str(sentence).replace(dep.text, str(trueOrFalsetemp(dep.label_,dep.text,ORG,PERSON,GPE_LOCATION,DATE,EVENT)))
			    	question_tf.append(temp_word)
			    	question_tf_key.append(dep.text)
			    	count+=1
			    	break
	    	
	count = 0
	keyword = ''
	choices = ''
	question_mcq = []
	question_mcq_choices = []
	question_mcq_key = []
	temp = ''
	questsent = ''
	fullsents = []
	tool = language_check.LanguageTool('en-US')
	for sentence in sentsWeight:
		
		text = sentence
		matches = tool.check(text)
		clntext = language_check.correct(text, matches)
		if len(clntext) < 20:
			continue
		sen = nlp(clntext)
		questsent = ''

		for ent in sen.noun_chunks:
			

			print(ent.text +" "+chunk.root.dep_+" "+ent.root.head.text)
			
			if ent.root.dep_ == 'nsubj':
				
				if temp == ent.root.head.text:
					pass
				else:
					temp = ''
					questsent += " "+ent.text +" "+ent.root.head.text
				temp = ent.root.head.text
				
			elif ent.root.dep_  == 'pobj':
				questsent += " "+ent.root.head.text +" "+ ent.text
				
			elif ent.root.dep_ == 'dobj':
				if temp == ent.root.head.text:
					pass
				else:
					temp = ''	
					questsent += " "+ent.root.head.text +" "+ ent.text
					
			elif ent.root.dep_ == 'appos':
					questsent += ent.root.head.text +" "+ ent.text
					
			elif ent.root.dep_ == 'attr':
					questsent += " "+ent.root.head.text +" "+ ent.text
					

		fullsents.append(questsent)
	for sents in fullsents:

		if len(sents) < 20:
			continue
		else:
			docs = nlp(sents)


		for parse in docs.ents:
			if parse.text not in NER:
				continue
			if parse.label_ == 'ORG':
				temp_word  = sents.replace(sents[0:parse.end_char], 'what ')

				if len(temp_word) < 5:
					break 

				question_mcq.append(temp_word)
				question_mcq_key.append(parse.text)
				question_mcq_choices.append(MCQ(parse.text,parse.label_,ORG,DATE,PERSON,GPE_LOCATION,EVENT))

			if  parse.label_ == 'PERSON':
				word = parse.text 
				end = parse.end_char
				if word[len(word)-2:len(word)] == "’s":

					temp_word  = sents.replace(sents[0:end-2], 'who')

				else:

					temp_word  = sents.replace(sents[0:parse.end_char], 'who ')

				if len(temp_word) < 3:

					break
				count +=1
				question_mcq.append(temp_word)
				question_mcq_key.append(parse.text)
				question_mcq_choices.append(MCQ(parse.text,parse.label_,ORG,DATE,PERSON,GPE_LOCATION,EVENT))


			if  parse.label_ == 'GPE':

				temp_word  = sents.replace(sents[0:parse.end_char], 'where ') 
				if len(temp_word) < 3:
					break
				question_mcq.append(temp_word)
				question_mcq_key.append(parse.text)
				question_mcq_choices.append(MCQ(parse.text,parse.label_,ORG,DATE,PERSON,GPE_LOCATION,EVENT))
				count +=1

			if  parse.label_ == 'EVENT':
				temp_word  = sents.replace(sents[0:parse.end_char], 'when ') 
				if len(temp_word) < 3:
					break 
				question_mcq.append(temp_word)
				question_mcq_key.append(parse.text)
				question_mcq_choices.append(MCQ(parse.text,parse.label_,ORG,DATE,PERSON,GPE_LOCATION,EVENT))
				count +=1
			if count == 20:
				break

	for i in question_tf:
		if len(i) <= 125:
			question_tf.remove(i)

	after = datetime.datetime.now()
	time = after - before

	question_mcq_choices_random = question_mcq_choices

	for i in range(len(question_mcq_choices_random)):
		for j in range(len(question_mcq_choices_random[i])):
			x = random.randrange(0,2)
			question_mcq_choices_random[i][x], question_mcq_choices_random[i][j] = question_mcq_choices_random[i][j], question_mcq_choices_random[i][x]

	y_true = NER
	y_pred = NER

	conf = confusion_matrix(y_true, y_pred, labels=NER)
	accuracy = getPrecisionRecall(conf)

	print(question_finb_key)
	for x in accuracy:
		print(x)


	ORG = UniqueItems(ORG)
	PERSON = UniqueItems(PERSON)
	GPE_LOCATION = UniqueItems(GPE_LOCATION)
	DATE = UniqueItems(DATE)
	EVENT = UniqueItems(EVENT)

	len_ORG = len(ORG)
	len_PERSON = len(PERSON)
	len_GPE_LOCATION = len(GPE_LOCATION)
	len_DATE = len(DATE)
	len_EVENT = len(EVENT)

	return render(request, 'genque/generate.html',
		{'NER': NER_Cleanse, 
		'filename': pdfFileObj.name, 
		'time': time, 
		'question_finb':question_finb,
		'question_finb_key': question_finb_key,
		'question_finb_key_random': question_finb_key_random,
		'question_tf':question_tf,
		'question_tf_key': question_tf_key,
		'question_mcq':question_mcq, 
		'question_mcq_key':question_mcq_key,
		'question_mcq_choices': question_mcq_choices,
		'ORG':ORG,
		'PERSON':PERSON,
		'GPE_LOCATION':GPE_LOCATION,
		'DATE':DATE,
		'EVENT':EVENT,
		'len_ORG':len_ORG,
		'len_PERSON':len_PERSON,
		'len_GPE_LOCATION':len_GPE_LOCATION,
		'len_DATE':len_DATE,
		'len_EVENT':len_EVENT,
		'accuracy':accuracy})


def download(request):
	finb_item = request.GET.getlist('finb_item')
	tf_item = request.GET.getlist('tf_item')
	mcq_item = request.GET.getlist('mcq_item')
	mcq_item_choices = request.GET.getlist('mcq_item_choices')
	choices = []

	for i in mcq_item_choices:
		choices.append(str(i).split(","))

	document = Document()

	document.add_heading("Name:___________________________________________________     Date:___________", level=3)
	document.add_heading("Section:________________________________________________      Score:__________", level=3)
	document.add_paragraph('')

	if not len(finb_item) == 0:
		document.add_heading("Fill in the Blanks",level = 3)
	for item in finb_item:
		document.add_paragraph(item,style = 'ListNumber')

	if not len(tf_item) == 0:
		document.add_heading("True or False",level = 3)
	for item in tf_item:
		document.add_paragraph(item,style = 'ListNumber')

	if not len(mcq_item) == 0:
		document.add_heading("Multiple Choices",level = 3)
	for item in mcq_item:
		document.add_paragraph(item[2:],style = 'ListNumber')
		for i in choices:
			if item[0] == i[0]:
				w = str(i[1:]).replace("'","")
				x = str(w).replace("[","")
				y = str(x).replace("]","")
				z = str(y).replace('"',"")
				choices_split = z.split(',')
				document.add_paragraph("a.)" + choices_split[0])
				document.add_paragraph("b.)" + choices_split[1])
				document.add_paragraph("c.)" + choices_split[2])



	doc_title = str(uuid.uuid4())+'.docx'
	document.save('genque/documents/' + doc_title)
	messages.info(request, 'Successfully downloaded the file!')

	filePath = 'genque/documents/' + doc_title
	fsock = open(filePath,"rb")
	response = HttpResponse(fsock, content_type='application/msword')
	response['Content-Disposition'] = 'attachment; filename='+doc_title
	return response

	# return HttpResponseRedirect('/')